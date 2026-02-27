"""
文档解析模块

支持多种文档格式的文本提取：
- PDF: 使用 PyMuPDF (fitz) 逐页提取文本
- DOCX: 使用 python-docx 提取段落文本
- TXT/MD: 自动检测编码后读取纯文本

采用策略模式：BaseParser 定义解析接口，各格式实现具体解析逻辑，
DocumentParser 作为统一入口自动选择对应的解析器。
"""

from pathlib import Path
from abc import ABC, abstractmethod
from collections import Counter
import re

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
import chardet


class BaseParser(ABC):
    """文档解析器基类，定义解析接口。"""

    @abstractmethod
    def parse(self, file_path: Path) -> str:
        """解析文档并返回提取的文本内容。"""
        pass

    @staticmethod
    def _squash_repeated_phrases(text: str) -> str:
        """
        压缩连续重复短语，缓解 OCR/模板导致的“字段名连写多次”问题。
        例如：'学院：学院：学院：' -> '学院：'
        """
        pattern = re.compile(r"(.{2,30}?)(?:\s*\1){1,}")
        previous = None
        while previous != text:
            previous = text
            text = pattern.sub(r"\1", text)
        return text

    @classmethod
    def _normalize_line(cls, line: str) -> str:
        """清理单行文本中的重复和多余空白。"""
        line = re.sub(r"\s+", " ", line).strip()
        if not line:
            return ""
        line = cls._squash_repeated_phrases(line)
        line = re.sub(r"([：:|])\1+", r"\1", line)
        return line

    @classmethod
    def _finalize_markdown(cls, blocks: list[str]) -> str:
        """去重并拼接 Markdown 片段。"""
        cleaned_blocks: list[str] = []
        prev_key = ""
        for block in blocks:
            text = block.strip()
            if not text:
                continue
            key = re.sub(r"\s+", " ", text)
            if key == prev_key:
                continue
            cleaned_blocks.append(text)
            prev_key = key
        return "\n\n".join(cleaned_blocks)

    @staticmethod
    def get_file_type(file_path: Path) -> str:
        """根据文件扩展名识别文档类型。"""
        suffix = file_path.suffix.lower()
        type_map = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "doc",
            ".txt": "txt",
            ".md": "markdown",
            ".markdown": "markdown",
        }
        return type_map.get(suffix, "unknown")


class PDFParser(BaseParser):
    """PDF 文档解析器，使用 PyMuPDF 逐页提取文本。"""

    def parse(self, file_path: Path) -> str:
        page_lines: list[list[str]] = []
        with fitz.open(file_path) as doc:
            for page in doc:
                # 纯文本模式：使用 blocks 提取并按坐标排序，尽量还原阅读顺序。
                blocks = page.get_text("blocks")
                blocks = sorted(blocks, key=lambda b: (b[1], b[0]))  # y0, x0

                lines = []
                for block in blocks:
                    if len(block) < 5:
                        continue
                    text = block[4]
                    if not isinstance(text, str) or not text.strip():
                        continue
                    for raw_line in text.splitlines():
                        line = self._normalize_line(raw_line)
                        if line:
                            lines.append(line)
                lines = [line for line in lines if line]
                if lines:
                    page_lines.append(lines)

        if not page_lines:
            return ""

        # 删除在大多数页面重复出现的短页眉/页脚文本。
        repeated_line_counter = Counter()
        for lines in page_lines:
            repeated_line_counter.update({line for line in lines if len(line) <= 40})

        threshold = max(2, int(len(page_lines) * 0.6))
        repeated_lines = {
            line for line, count in repeated_line_counter.items() if count >= threshold
        }

        md_blocks: list[str] = []
        for idx, lines in enumerate(page_lines, 1):
            filtered = [line for line in lines if line not in repeated_lines]
            if not filtered:
                continue
            md_blocks.append(f"## 第{idx}页\n\n" + "\n\n".join(filtered))

        return self._finalize_markdown(md_blocks)


class DocxParser(BaseParser):
    """Word 文档解析器，使用 python-docx 提取段落文本。"""

    def parse(self, file_path: Path) -> str:
        # 主流程：先用 mammoth 将 docx 转 markdown，结构化效果通常更好。
        markdown = self._parse_with_mammoth(file_path)
        if markdown:
            return markdown

        # 回退流程：mammoth 不可用或失败时，使用 python-docx 提取。
        return self._parse_with_python_docx(file_path)

    def _parse_with_mammoth(self, file_path: Path) -> str:
        """使用 mammoth 将 docx 转为 markdown。失败时返回空字符串。"""
        try:
            import mammoth
        except Exception:
            return ""

        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_markdown(docx_file)
            md = result.value or ""
            if not md.strip():
                return ""
            return self._normalize_markdown(md)
        except Exception:
            return ""

    def _normalize_markdown(self, markdown: str) -> str:
        """清洗 mammoth 输出的 markdown，去掉重复空行与重复文本。"""
        blocks: list[str] = []
        for raw_line in markdown.splitlines():
            stripped = raw_line.rstrip()
            if not stripped:
                blocks.append("")
                continue

            # markdown 语法前缀（标题/列表/引用）保留，仅清洗正文内容。
            m = re.match(r"^([#>\-\*\d\.\)\s]+)(.*)$", stripped)
            if m:
                prefix, content = m.group(1), m.group(2)
                content = self._normalize_line(content)
                line = f"{prefix}{content}".rstrip() if content else prefix.strip()
            else:
                line = self._normalize_line(stripped)
            if line:
                blocks.append(line)

        # 合并连续空行
        collapsed: list[str] = []
        prev_empty = True
        for line in blocks:
            is_empty = (line == "")
            if is_empty and prev_empty:
                continue
            collapsed.append(line)
            prev_empty = is_empty

        # 对非空块做相邻去重
        joined = "\n".join(collapsed)
        pieces = [p for p in joined.split("\n\n") if p.strip()]
        return self._finalize_markdown(pieces)

    def _parse_with_python_docx(self, file_path: Path) -> str:
        """回退解析：使用 python-docx 提取并转为 markdown。"""
        doc = DocxDocument(file_path)
        md_blocks: list[str] = []

        # 按正文原始顺序提取段落和表格并转为 Markdown。
        for child in doc.element.body.iterchildren():
            tag = child.tag
            if tag.endswith("}p"):
                para_md = self._paragraph_to_markdown(child, doc)
                if para_md:
                    md_blocks.append(para_md)
            elif tag.endswith("}tbl"):
                table = Table(child, doc)
                table_md = self._table_to_markdown(table)
                if table_md:
                    md_blocks.append(table_md)

        # 补充页眉/页脚中的文字（常见于公文模板）。
        header_footer_seen = set()
        for section in doc.sections:
            for para in section.header.paragraphs:
                text = self._normalize_line(para.text)
                if text and text not in header_footer_seen:
                    md_blocks.append(text)
                    header_footer_seen.add(text)
            for para in section.footer.paragraphs:
                text = self._normalize_line(para.text)
                if text and text not in header_footer_seen:
                    md_blocks.append(text)
                    header_footer_seen.add(text)

        # 兜底：补充文本框中的内容，避免复杂模板漏字。
        for text in self._extract_textbox_lines(doc):
            md_blocks.append(text)

        return self._finalize_markdown(md_blocks)

    def _paragraph_to_markdown(self, p_node, doc: DocxDocument) -> str:
        """将段落节点转为 Markdown（标题/列表/普通段落）。"""
        para = Paragraph(p_node, doc)
        text = self._normalize_line(para.text or "")
        if not text:
            return ""

        # 标题样式
        style_name = (para.style.name or "").lower() if para.style else ""
        if "heading" in style_name:
            level_match = re.search(r"(\d+)", style_name)
            level = min(6, int(level_match.group(1))) if level_match else 2
            return f"{'#' * level} {text}"

        # 编号/项目符号列表
        has_num = bool(p_node.xpath(".//*[local-name()='numPr']"))
        if has_num:
            return f"- {text}"

        return text

    def _extract_textbox_lines(self, doc: DocxDocument) -> list[str]:
        """提取文本框内容（w:txbxContent）。"""
        lines: list[str] = []
        for node in doc.element.xpath(".//*[local-name()='txbxContent']//*[local-name()='t']"):
            text = self._normalize_line(node.text or "")
            if text:
                lines.append(text)
        return lines

    def _table_to_markdown(self, table: Table) -> str:
        """将 Word 表格转换为 Markdown，优先保留非空信息，减少 | | | 噪声。"""
        rows: list[list[str]] = []
        for row in table.rows:
            cells: list[str] = []
            prev_cell = ""
            for cell in row.cells:
                # 用底层 XML 提取 cell 全量文本，避免 cell.text 在复杂结构下漏值。
                texts = cell._tc.xpath(".//*[local-name()='t']/text()")
                cell_text = self._normalize_line(" ".join(texts))
                if not cell_text:
                    continue
                # 合并单元格时 python-docx 可能在相邻单元格返回重复值，连续去重一次。
                if cell_text == prev_cell:
                    continue
                cells.append(cell_text)
                prev_cell = cell_text
            if cells:
                rows.append(cells)

        if not rows:
            return ""

        # 行列相对规整时输出标准 Markdown 表格；否则按列表输出，避免空列噪声。
        col_lens = {len(r) for r in rows}
        if len(col_lens) == 1 and next(iter(col_lens)) >= 2:
            col_count = next(iter(col_lens))
            header = rows[0]
            body_rows = rows[1:] if len(rows) > 1 else []
            lines = [
                "| " + " | ".join(header) + " |",
                "| " + " | ".join(["---"] * col_count) + " |",
            ]
            for row in body_rows:
                lines.append("| " + " | ".join(row) + " |")
            return "\n".join(lines)

        return "\n".join(f"- {' | '.join(row)}" for row in rows)


class TextParser(BaseParser):
    """纯文本/Markdown 解析器，自动检测文件编码。"""

    def parse(self, file_path: Path) -> str:
        with open(file_path, "rb") as f:
            raw_data = f.read()

        # 使用 chardet 自动检测文件编码
        detected = chardet.detect(raw_data)
        encoding = detected.get("encoding", "utf-8") or "utf-8"

        try:
            return raw_data.decode(encoding)
        except UnicodeDecodeError:
            # 检测失败时回退到 UTF-8，忽略无法解码的字符
            return raw_data.decode("utf-8", errors="ignore")


class DocumentParser:
    """
    统一文档解析入口。

    根据文件扩展名自动选择对应的解析器，
    返回提取的文本内容和文件类型。
    """

    # 文件类型 → 解析器实例的映射表
    _parsers: dict[str, BaseParser] = {
        "pdf": PDFParser(),
        "docx": DocxParser(),
        "txt": TextParser(),
        "markdown": TextParser(),
    }

    @classmethod
    def parse(cls, file_path: Path) -> tuple[str, str]:
        """
        解析文档文件。

        Args:
            file_path: 文档文件路径

        Returns:
            (提取的文本内容, 文件类型)

        Raises:
            ValueError: 不支持的文件格式
        """
        file_type = BaseParser.get_file_type(file_path)

        if file_type == "unknown":
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

        if file_type == "doc":
            raise ValueError("Legacy .doc format not supported. Please convert to .docx")

        parser = cls._parsers.get(file_type)
        if not parser:
            raise ValueError(f"No parser available for file type: {file_type}")

        content = parser.parse(file_path)
        return content, file_type

    @classmethod
    def supported_extensions(cls) -> list[str]:
        """返回支持的文件扩展名列表。"""
        return [".pdf", ".docx", ".txt", ".md", ".markdown"]
